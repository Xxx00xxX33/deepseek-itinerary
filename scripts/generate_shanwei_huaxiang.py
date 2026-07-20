#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
华宇报价（汕尾回乡团20人） 更改.doc -> Voyage 风格专业行程 DOCX

严格遵守：
- 报价为唯一事实来源
- 不修改原有行程安排
- 行程中（Day条、正文）不出现任何价格数字
- 不编造景点、服务、赠送
- 输出简体中文 .docx
- 克隆已知 Voyage 模板保证 Word 可打开

参考: generate_chaoshan6_voyage_style.py + SKILL.md
"""

from __future__ import annotations

import argparse
import os
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 模板（保持公司抬头、页边距、主题）
TEMPLATE = r'C:\Users\DELL\Downloads\12D11N_东北草原深度游_2026年6月.docx'


def get_output_paths(source_doc: Path):
    """输出到源目录：与输入 .doc 同目录，中文主名 + ASCII 副本。"""
    out_dir = source_doc.parent
    primary = out_dir / '华宇汕尾回乡团（新加坡往返）行程.docx'
    ascii_copy = out_dir / 'Shanwei_Hometown_Tour_Itinerary.docx'
    return [primary, ascii_copy]

# 颜色（Voyage / 东北草原风格）
C_TITLE = RGBColor(0x1A, 0x6B, 0x5E)
C_SUB = RGBColor(0x55, 0x55, 0x55)
C_ORANGE = RGBColor(0xC4, 0x59, 0x11)
C_DAY_BG = '2E8B7A'
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY = RGBColor(0x00, 0x00, 0x00)
C_HOTEL = RGBColor(0x1A, 0x6B, 0x5E)

FONT_CN = 'SimSun'
FONT_DAY = 'Arial Unicode MS'


def set_run_font(run, name=FONT_CN, size_pt=11, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is True:
        run.font.bold = True
        if rPr.find(qn('w:bCs')) is None:
            rPr.append(OxmlElement('w:bCs'))
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def set_spacing(paragraph, before=None, after=None, left=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if left is not None:
        pf.left_indent = Pt(left)
    if line is not None:
        pf.line_spacing = line


def clear_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def add_para(doc, text, *, size=11, bold=False, color=C_BODY, font=FONT_CN,
             align=None, before=0, after=4, left=None, justify=False,
             shade=None, line=1.15):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align is not None:
        p.alignment = align
    set_spacing(p, before=before, after=after, left=left, line=line)
    if shade:
        set_paragraph_shading(p, shade)
    if text:
        run = p.add_run(text)
        set_run_font(run, name=font, size_pt=size, bold=bold, color=color)
    return p


def add_sight_para(doc, name: str, desc: str):
    """【景点名】加粗深绿 + 说明正文，便于扫读。不重复概述段落。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=2, after=3, line=1.2)
    r1 = p.add_run(f'【{name}】')
    set_run_font(r1, size_pt=10.5, bold=True, color=C_TITLE)
    r2 = p.add_run(desc)
    set_run_font(r2, size_pt=10.5, color=C_BODY)
    return p


def add_label_value(doc, label: str, value: str, *, before=2, after=2):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after, line=1.15)
    r1 = p.add_run(label)
    set_run_font(r1, size_pt=10.5, bold=True, color=C_TITLE)
    r2 = p.add_run(value)
    set_run_font(r2, size_pt=10.5, color=C_BODY)
    return p


# ==================== 仅来自报价的事实（价格数字仅出现在底部报价与说明区） ====================

TITLE = '汕尾回乡团（新加坡往返）'
SUBTITLE = '新加坡·深圳·汕尾·海丰·陆丰·陆河'

NOTE_NO_SHOP = '***全程不进店'

GIFTS = '每天每人矿泉水两瓶'

# 严格按报价逐日原文（去除价格数字）。概述段落紧贴报价。
# 景点介绍约60-70中文字，面向非中国游客，客观且更有吸引力，避免民族或宏大叙事，仅基于地点名称与报价上下文。
DAYS = [
    {
        'day': 'Day 1',
        'route': '新加坡/深圳',
        'meals': '抵 达',
        'summary': '接机，入住酒店，自由逛东门步行街。',
        'sights': [
            ('东门步行街', '东门步行街是深圳热闹繁华的商业步行街，街道两旁林立着各式商店和餐厅，游客可以随意逛街购物，感受这座城市夜晚的活力与繁忙氛围。'),
        ],
        'hotel': '中泰来或同级四星',
    },
    {
        'day': 'Day 2',
        'route': '深圳/汕尾（2.5H）',
        'meals': '早：酒店｜午：粤菜风味｜晚：海鲜晚宴',
        'summary': '早餐后前往汕尾，善美广场、品清湖夜景漫步。',
        'sights': [
            ('善美广场', '善美广场坐落在汕尾市区中心，是当地居民日常休闲的公共空间，广场周边商业设施完善，游客可在此悠闲散步，近距离观察和体验本地人的日常生活场景。'),
            ('品清湖', '品清湖是汕尾市区一处开阔宁静的天然湖泊，湖面宽广平静如镜，夜晚华灯映照下水面波光粼粼，游客沿着湖畔步道漫步，能充分欣赏那迷人的湖光夜色。'),
        ],
        'hotel': '维也纳国际豪华房（城区店，四星）',
    },
    {
        'day': 'Day 3',
        'route': '汕尾/海丰/汕尾（0.5+0.5H）',
        'meals': '午：海丰风味｜晚：自理',
        'summary': '参观海丰红宫红场旧址纪念馆、彭湃烈士故居，午餐后前往凤山民俗文化旅游区（非遗展演已经没有了，改凤山），逛二马路。',
        'sights': [
            ('海丰红宫红场旧址纪念馆', '海丰红宫红场旧址纪念馆由一组保存完好的历史建筑群构成，馆内通过丰富展览详细介绍当地历史人物彭湃的事迹及其故居，游客可在这里深入了解这段往昔的历史。'),
            ('彭湃烈士故居', '彭湃烈士故居是一座保存完整的传统民居建筑，内部陈列着大量历史文物和生活用品，游客参观时能够真切感受到当地过去普通居民的居住环境和生活方式。'),
            ('凤山民俗文化旅游区', '凤山民俗文化旅游区位于海丰县城，区内设有民俗展示和商业街区，游客可沿着热闹的二马路缓缓行走，浏览各类当地店铺，感受浓郁的市井生活气息。'),
        ],
        'hotel': '维也纳国际豪华房（城区店，四星）',
    },
    {
        'day': 'Day 4',
        'route': '汕尾/陆丰/汕尾（1.5+1.5H）',
        'meals': '午：陆丰风味｜晚：牛肉火锅',
        'summary': '前往陆丰，游览玄武山、定光禅寺、金厢银滩。',
        'sights': [
            ('玄武山', '玄武山是陆丰地区著名的山地自然景区，山体起伏绵延，山顶建有古寺庙宇，游客沿着蜿蜒山径登山而上，可一边欣赏沿途茂密林木，一边俯瞰周围壮丽的自然山林景色。'),
            ('定光禅寺', '定光禅寺是陆丰一座历史悠久的佛教寺院，寺内殿堂建筑古朴庄重，院落布局清幽整洁，环境宁静祥和，游客在此可以悠闲参观，感受传统佛教文化的宁静氛围。'),
            ('金厢银滩', '金厢银滩位于陆丰沿海地带，是面积开阔的优质沙滩景区，沙质洁白细腻，海水清澈透明，游客可以在宽阔的沙滩上自由散步，感受海风拂面和海浪轻拍的惬意。'),
        ],
        'hotel': '维也纳国际豪华房（城区店，四星）',
    },
    {
        'day': 'Day 5',
        'route': '汕尾/陆河/汕尾（1.5+1.5H）',
        'meals': '午：客家风味（客家酿豆腐、擂茶、咸鸡）｜晚：全牛宴',
        'summary': '前往陆河螺洞世外梅园景区、九厅十八井客家古民居，后返回汕尾与侨联座谈。',
        'sights': [
            ('螺洞世外梅园景区', '螺洞世外梅园景区地处陆河，拥有大片连绵的梅花种植园和秀丽田园风光，游客漫步在园中曲径上，能够欣赏到季节花海与乡村田野交织的自然美景，享受宁静的田园情趣。'),
            ('九厅十八井客家古民居', '九厅十八井客家古民居是陆河地区典型的传统围屋式建筑群落，由多个厅堂和水井巧妙组成，整体布局严谨对称，游客进入参观时可以近距离感受这种独特民居的建筑结构和历史风貌。'),
        ],
        'hotel': '维也纳国际豪华房（城区店，四星）',
    },
    {
        'day': 'Day 6',
        'route': '汕尾/深圳/新加坡（3H）',
        'meals': '午：全蚝宴',
        'summary': '参观晨洲蚝乡，自由购物，午餐后前往深圳，送机。',
        'sights': [
            ('晨洲蚝乡', '晨洲蚝乡是汕尾沿海著名的牡蛎养殖村落，村子紧邻大海，游客可以走进蚝田区域近距离观察养殖过程，了解当地渔民世代相传的水产养殖技艺和海边特有的产业生活景象。'),
        ],
        'hotel': None,  # 返程日无宿
    },
]


def build(outputs: list[Path]):
    if not os.path.exists(TEMPLATE):
        raise SystemExit(f'TEMPLATE not found: {TEMPLATE}')

    doc = Document(TEMPLATE)
    clear_body(doc)

    # 标题区
    add_para(doc, TITLE, size=18, bold=True, color=C_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=3, line=1.2)
    add_para(doc, SUBTITLE, size=11, bold=True, color=C_SUB,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line=1.15)
    add_para(doc, NOTE_NO_SHOP, size=12, bold=True, color=C_ORANGE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8, line=1.1)

    # 仅报价明确有的赠送（无独家安排 / 各地美食 区块，因报价未提供）
    add_label_value(doc, '赠送：', GIFTS + '。', before=1, after=8)

    # 每日行程
    for day in DAYS:
        # Day 条：无多余空格填充，使用 2-4 空格分隔；餐食名称不带价格
        header = f'{day["day"]}  {day["route"]}    （{day["meals"]}）'
        add_para(doc, header, size=11, bold=True, color=C_WHITE, font=FONT_DAY,
                 before=10, after=4, left=2, shade=C_DAY_BG, line=1.1)

        # 当日概述（两端对齐，紧贴报价原文）
        add_para(doc, day['summary'], size=10.5, color=C_BODY, justify=True,
                 before=3, after=3, line=1.25)

        # 景点扫读行（若有）
        for name, desc in day.get('sights', []):
            add_sight_para(doc, name, desc)

        # 宿（返程日无）
        if day['hotel']:
            add_para(doc, f'宿：{day["hotel"]}', size=10.5, bold=True, color=C_HOTEL,
                     before=4, after=2, left=2, line=1.1)

    # 报价与说明单独一页（与行程正文分开）
    doc.add_page_break()
    add_para(doc, '报价与说明', size=12, bold=True, color=C_TITLE,
             before=6, after=6, line=1.1)

    add_para(
        doc,
        '报价：RMB2230/人，单间差RMB750，20人以上报价，16免1，用车37座。',
        size=10.5, bold=True, color=C_BODY, before=2, after=4, line=1.2,
    )

    add_label_value(doc, '购物：', '全程不进店', before=1, after=1)
    add_label_value(doc, '小费：', 'RMB30/人天', before=0, after=1)
    add_label_value(doc, '赠送：', '每天每人矿泉水两瓶', before=0, after=3)

    add_para(
        doc,
        '精选酒店：首晚 中泰来或同级四星；其余晚 维也纳国际豪华房（城区店，四星）。',
        size=10.5, color=C_BODY, before=2, after=4, line=1.2,
    )

    add_para(
        doc,
        '膳食（据报价）：抵 达日 接机入住；深圳至汕尾日 早酒店、午粤菜风味、晚海鲜晚宴；'
        '汕尾海丰往返日 午海丰风味、晚自理；汕尾陆丰往返日 午陆丰风味、晚牛肉火锅；'
        '汕尾陆河往返日 午客家风味（客家酿豆腐、擂茶、咸鸡）、晚全牛宴；返程日 午全蚝宴。',
        size=10.5, color=C_BODY, before=1, after=6, line=1.25,
    )

    add_para(doc, '联系方式', size=12, bold=True, color=C_TITLE,
             before=4, after=3, line=1.1)
    add_para(doc, '联系人：姚文跃（13600926388） 李珏（13950127910）',
             size=10.5, color=C_BODY, before=0, after=1, line=1.15)
    add_para(doc, '电话：0086-592-8121184  8121185  2031766  2036977',
             size=10.5, color=C_BODY, before=0, after=1, line=1.15)
    add_para(doc, '传真：0086-592-8121195  2027855',
             size=10.5, color=C_BODY, before=0, after=1, line=1.15)
    add_para(doc, 'Email：yao19760126@aliyun.com',
             size=10.5, color=C_BODY, before=0, after=4, line=1.15)

    # 保存主文件 + ASCII 副本（输出到源目录）
    primary = outputs[0]
    doc.save(primary)
    print('Saved:', primary, os.path.getsize(primary))
    for other in outputs[1:]:
        shutil.copy2(primary, other)
        print('Copied:', other)


def main():
    ap = argparse.ArgumentParser(description='Generate Voyage itinerary DOCX next to source quote .doc')
    ap.add_argument('quote', nargs='?', default=None,
                    help='Path to the source .doc quote (output will be written beside it)')
    args = ap.parse_args()

    if args.quote:
        source = Path(args.quote).resolve()
    else:
        # Fallback to the known quote for this product
        source = Path(r'C:\Users\DELL\Downloads\华宇报价（汕尾回乡团20人） 更改.doc')

    if not source.exists():
        raise SystemExit(f'Source quote not found: {source}')

    outputs = get_output_paths(source)
    build(outputs)


if __name__ == '__main__':
    main()
