#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Worked example for deepseek-itinerary skill (Voyage paragraph style).

潮汕6天：克隆东北草原深度游 DOCX 作壳，行程/价格严格按华宇报价，
优化空格与可读性。For new products, copy patterns into generate.py.

See ../SKILL.md and ../references/fact-check.md.
"""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

TEMPLATE = r'C:\Users\DELL\Downloads\12D11N_东北草原深度游_2026年6月.docx'
OUTPUTS = [
    r'C:\Users\DELL\Downloads\跟着阿嬷游潮汕_厦门进出6天行程.docx',
    r'C:\Users\DELL\Downloads\Chaoshan_Xiamen_6D_Itinerary.docx',
    r'C:\Users\DELL\Downloads\chaoshan6_itinerary.docx',
]

C_TITLE = RGBColor(0x1A, 0x6B, 0x5E)
C_SUB = RGBColor(0x55, 0x55, 0x55)
C_ORANGE = RGBColor(0xC4, 0x59, 0x11)
C_HL = RGBColor(0x41, 0x9D, 0x99)
C_FOOD = RGBColor(0x2E, 0x8B, 0x7A)
C_DAY_BG = '2E8B7A'
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY = RGBColor(0x00, 0x00, 0x00)
C_HOTEL = RGBColor(0x1A, 0x6B, 0x5E)

FONT_CN = 'SimSun'
FONT_DAY = 'Arial Unicode MS'


def set_run_font(run, name=FONT_CN, size_pt=11, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is True:
        run.font.bold = True
        if rPr.find(qn('w:bCs')) is None:
            rPr.append(OxmlElement('w:bCs'))
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def set_spacing(paragraph, before=None, after=None, left=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if left is not None:
        pf.left_indent = Pt(left)
    if line is not None:
        pf.line_spacing = line


def clear_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def add_para(doc, text, *, size=11, bold=False, color=C_BODY, font=FONT_CN,
             align=None, before=0, after=4, left=None, justify=False,
             shade=None, line=1.15):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align is not None:
        p.alignment = align
    set_spacing(p, before=before, after=after, left=left, line=line)
    if shade:
        set_paragraph_shading(p, shade)
    if text:
        run = p.add_run(text)
        set_run_font(run, name=font, size_pt=size, bold=bold, color=color)
    return p


def add_sight_para(doc, name: str, desc: str):
    """【景点名】加粗 + 说明正文，便于扫读。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=2, after=3, line=1.2)
    r1 = p.add_run(f'【{name}】')
    set_run_font(r1, size_pt=10.5, bold=True, color=C_TITLE)
    r2 = p.add_run(desc)
    set_run_font(r2, size_pt=10.5, color=C_BODY)
    return p


def add_label_value(doc, label: str, value: str, *, before=2, after=2):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after, line=1.15)
    r1 = p.add_run(label)
    set_run_font(r1, size_pt=10.5, bold=True, color=C_TITLE)
    r2 = p.add_run(value)
    set_run_font(r2, size_pt=10.5, color=C_BODY)
    return p


# ---------- 内容（报价事实不变；去掉多余空格、理顺层次） ----------

TITLE = '6天跟着阿嬷游潮汕·舌尖上的福建/潮汕美食团'
SUBTITLE = '厦门·揭阳·永定土楼·潮州·汕头·澄海·南澳岛'

NOTE_NO_SHOP = '***全程不进店'

HIGHLIGHT_LINES = [
    '跟着阿嬷游潮汕，走进《给阿嬷的情书》故里，领略潮汕华侨文化',
    '漫步千年潮州古城，百年茶馆品茗听潮剧；欣赏非遗文化潮汕英歌舞',
    '手工DIY红桃粿；广东十大最美古村—龙湖古寨',
    '广东最美岛屿、东方夏威夷—南澳岛',
    '畅游海上花园鼓浪屿、永定土楼两大世界文化遗产',
]

FOOD = (
    '龙虾鲍鱼红酒晚宴、北京烤鸭、自助海鲜烤肉火锅、潮式十八碟（潮汕小吃）、'
    '汕头全鹅宴、客家风味、梅菜扣肉+盐酒鸡、揭阳小吃、南澳海鲜'
)

GIFTS = (
    '每人每天矿泉水两瓶；行程内另含：赠送温泉门票（享受温泉SPA，请自备泳衣）、'
    '赠送品尝客家米酒、赠送品尝南枝同款“无米粿”'
)

DAYS = [
    {
        'day': 'Day 1',
        'route': '新加坡→厦门',
        'meals': '晚：龙虾鲍鱼红酒晚宴',
        'summary': (
            '搭机前往福建最美滨海城市—厦门，参观环岛路黄金海岸、演武大桥海上观景平台，'
            '远眺厦门地标—世贸双子塔，打卡网红景点—沙坡尾。'
            '晚餐：龙虾鲍鱼红酒晚宴150。'
        ),
        'sights': [
            ('环岛路黄金海岸、演武大桥海上观景平台',
             '参观环岛路黄金海岸、演武大桥海上观景平台。'),
            ('世贸双子塔、沙坡尾',
             '远眺厦门地标—世贸双子塔，打卡网红景点—沙坡尾。本日报价膳食仅列晚餐。'),
        ],
        'hotel': '丽华酒店（五星）',
    },
    {
        'day': 'Day 2',
        'route': '厦门→永定（3H）',
        'meals': '早：酒店｜午：北京烤鸭｜晚：客家料理',
        'summary': (
            '早餐后参观世界文化遗产、海上花园—鼓浪屿（龙头路商业街、万国建筑博览、邂逅最美转角），'
            '前往永定，游览客家土楼中最美丽、最原生态的土楼群—初溪土楼群（含电瓶车），'
            '参观世界文化遗产—集庆楼、善庆楼、绳庆楼；'
            '入住当地最好温泉酒店，尽情享受泡汤乐趣（赠送温泉门票，享受温泉SPA，请自备泳衣）。'
            '晚餐品尝正宗客家料理（赠送品尝客家米酒）。'
            '午餐：北京烤鸭60；晚餐：客家料理60。'
        ),
        'sights': [
            ('鼓浪屿',
             '世界文化遗产、海上花园。含龙头路商业街、万国建筑博览、邂逅最美转角。'),
            ('初溪土楼群（含电瓶车）·集庆楼、善庆楼、绳庆楼',
             '客家土楼中最美丽、最原生态的土楼群；参观世界文化遗产—集庆楼、善庆楼、绳庆楼。'),
            ('天子温泉',
             '赠送温泉门票，享受温泉SPA，请自备泳衣；晚餐品尝正宗客家料理（赠送品尝客家米酒）。'),
        ],
        'hotel': '天子温泉（五星）',
    },
    {
        'day': 'Day 3',
        'route': '永定→揭阳→潮州（3H+1H）',
        'meals': '早：酒店｜午：揭阳风味｜晚：潮式十八碟',
        'summary': (
            '早餐后前往揭阳，打卡《给阿嬷的情书》取景地—揭阳西淇村，'
            '感受古村田园诗意与乡愁，一纸侨批，半世守望，'
            '《石板桥》承载着阿嬷的牵挂与等待，赠送品尝南枝同款“无米粿”；'
            '随后打卡揭阳古城西马路历史文化名街（电影唐人街、暹罗客栈取景地），'
            '南洋风情的骑楼错落有致，青石板路藏岁月最深记忆；'
            '后往历史古都潮州，游览广东十大最美、千年古村—龙湖古寨（还原电影90%的泰国街景）。'
            '午餐：揭阳风味60；晚餐：潮式十八碟80（手工DIY红桃粿）。'
        ),
        'sights': [
            ('揭阳西淇村（《给阿嬷的情书》取景地）',
             '古村田园与侨批故事；《石板桥》；赠送品尝南枝同款“无米粿”。'),
            ('揭阳古城西马路历史文化名街',
             '电影唐人街、暹罗客栈取景地；南洋风情骑楼、青石板路。'),
            ('龙湖古寨',
             '广东十大最美古村、千年古村（还原电影90%的泰国街景）。'),
        ],
        'hotel': '愉羽酒店（准五星）',
    },
    {
        'day': 'Day 4',
        'route': '潮州→澄海→汕头→潮州（1H+1H+1H）',
        'meals': '早：酒店｜午：汕头全鹅宴｜晚：自理潮汕小吃',
        'summary': (
            '早餐后前往澄海，打卡《给阿嬷的情书》取景地—岭南第一侨宅—陈慈黉故居'
            '（欣赏非遗文化潮汕英歌舞）；前往美食之乡—汕头，参观小公园老城骑楼建筑群'
            '（海平路77号还原片中的“暹南电影院”、裕丰银信局（万安街3号还原剧中寄侨批的名场面）'
            '+侨批文物馆（逢周一闭馆））；后返潮州，游览潮州古城（含电瓶车）：'
            '打卡西湖（湖心亭）+泰佛殿，广济门古城楼、牌坊街、百年茶馆品茗听潮剧、'
            '漫步滨江长廊，远眺湘子桥；晚上欣赏古城灯光秀（雨天取消）。'
            '午餐：汕头全鹅宴60；晚餐：自理潮汕小吃。'
        ),
        'sights': [
            ('陈慈黉故居（岭南第一侨宅）',
             '《给阿嬷的情书》取景地；欣赏非遗文化潮汕英歌舞。'),
            ('小公园老城骑楼·侨批文物馆',
             '海平路77号“暹南电影院”；万安街3号裕丰银信局；侨批文物馆（逢周一闭馆）。'),
            ('潮州古城（含电瓶车）',
             '西湖（湖心亭）+泰佛殿、广济门古城楼、牌坊街、百年茶馆品茗听潮剧、'
             '滨江长廊、远眺湘子桥；古城灯光秀（雨天取消）。'),
        ],
        'hotel': '愉羽酒店（准五星）',
    },
    {
        'day': 'Day 5',
        'route': '潮州→南澳→厦门（1.5H+3.5H）',
        'meals': '早：酒店｜午：南澳小海鲜｜晚：自助海鲜烤肉火锅',
        'summary': (
            '早餐后前往“广东最美岛屿”—南澳岛，途经南澳大桥，'
            '车游新南澳外滩—前江湾海滨路，打卡网红景点—启航广场，'
            '参观自然之门—北回归线标志塔、东方夏威夷—青澳湾，后往厦门。'
            '午餐：南澳小海鲜60；晚餐：自助海鲜烤肉火锅100。'
        ),
        'sights': [
            ('南澳岛·南澳大桥·前江湾·启航广场',
             '途经南澳大桥；车游新南澳外滩—前江湾海滨路；打卡网红景点—启航广场。'),
            ('北回归线标志塔（自然之门）、青澳湾（东方夏威夷）',
             '参观自然之门—北回归线标志塔、东方夏威夷—青澳湾，后往厦门。'),
        ],
        'hotel': '丽华酒店（五星）',
    },
    {
        'day': 'Day 6',
        'route': '厦门→新加坡（返程）',
        'meals': '早：酒店',
        'summary': '早餐后自由活动至集合时间，前往机场搭机返回温暖家园。',
        'sights': [],
        'hotel': None,
    },
]


def build():
    doc = Document(TEMPLATE)
    clear_body(doc)

    # —— 封面信息 ——
    add_para(doc, TITLE, size=18, bold=True, color=C_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=3, line=1.2)
    add_para(doc, SUBTITLE, size=11, bold=True, color=C_SUB,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line=1.15)
    add_para(doc, NOTE_NO_SHOP, size=12, bold=True, color=C_ORANGE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8, line=1.1)

    # —— 导语块：标签+内容，避免长句糊成一团 ——
    add_label_value(doc, '独家安排：', '；'.join(HIGHLIGHT_LINES) + '。', before=2, after=4)
    add_label_value(doc, '各地美食：', FOOD + '。', before=1, after=3)
    add_label_value(doc, '赠送：', GIFTS + '。', before=1, after=8)

    # —— 每日行程 ——
    for day in DAYS:
        # Day 条：左 Day+路线，右餐食，去掉大段空格填充
        header = f'{day["day"]}  {day["route"]}    （{day["meals"]}）'
        add_para(doc, header, size=11, bold=True, color=C_WHITE, font=FONT_DAY,
                 before=10, after=4, left=2, shade=C_DAY_BG, line=1.1)

        add_para(doc, day['summary'], size=10.5, color=C_BODY, justify=True,
                 before=3, after=3, line=1.25)

        for name, desc in day['sights']:
            add_sight_para(doc, name, desc)

        if day['hotel']:
            add_para(doc, f'宿：{day["hotel"]}', size=10.5, bold=True, color=C_HOTEL,
                     before=4, after=2, left=2, line=1.1)

    # —— 报价与备注 ——
    add_para(doc, '报价与说明', size=12, bold=True, color=C_TITLE,
             before=12, after=6, line=1.1)

    add_para(
        doc,
        '报价：RMB2480/人，单间差RMB800，16人以上报价，16免1'
        '（领队若用单间需补单间差），小孩占床和大人同价，不占床按60%。',
        size=10.5, bold=True, color=C_BODY, before=2, after=4, line=1.2,
    )

    add_label_value(doc, '购物：', '全程不进店', before=1, after=1)
    add_label_value(doc, '小费：', 'RMB30/人天', before=0, after=1)
    add_label_value(doc, '赠送：', '每人每天矿泉水两瓶', before=0, after=3)

    add_para(
        doc,
        '备注：此报价有效期为2026年6月1日至2027年12月31日，'
        '但需避开1-5/5“五一节”、1-7/10“国庆节”、春节期间。',
        size=10.5, color=C_BODY, before=2, after=4, line=1.2,
    )

    add_para(
        doc,
        '精选酒店：全程当地5*酒店，提升1晚超五星温泉酒店（享受温泉SPA）。'
        '行程所列：D1/D5丽华酒店（五星）；D2天子温泉（五星）；D3/D4愉羽酒店（准五星）。',
        size=10.5, color=C_BODY, before=1, after=4, line=1.2,
    )

    add_para(
        doc,
        '膳食（据报价）：'
        'D1晚餐龙虾鲍鱼红酒晚宴150；'
        'D2三餐（午北京烤鸭60、晚客家料理60）；'
        'D3三餐（午揭阳风味60、晚潮式十八碟80含手工DIY红桃粿）；'
        'D4早午（午汕头全鹅宴60）、晚自理潮汕小吃；'
        'D5三餐（午南澳小海鲜60、晚自助海鲜烤肉火锅100）；'
        'D6早餐。',
        size=10.5, color=C_BODY, before=1, after=6, line=1.25,
    )

    add_para(doc, '联系方式', size=12, bold=True, color=C_TITLE,
             before=4, after=3, line=1.1)
    add_para(doc, '联系人：姚文跃(13600926388)　李珏(13950127910)',
             size=10.5, color=C_BODY, before=0, after=1, line=1.15)
    add_para(doc, '电话：0592-8121184　8121185　传真：0592-8121195',
             size=10.5, color=C_BODY, before=0, after=1, line=1.15)
    add_para(doc, 'Skype：xiaoyao_1976　Email：yao19760126@aliyun.com　QQ：499756025',
             size=10.5, color=C_BODY, before=0, after=4, line=1.15)

    first = OUTPUTS[0]
    doc.save(first)
    print('Saved:', first, os.path.getsize(first))
    for other in OUTPUTS[1:]:
        shutil.copy2(first, other)
        print('Copied:', other)


if __name__ == '__main__':
    build()
