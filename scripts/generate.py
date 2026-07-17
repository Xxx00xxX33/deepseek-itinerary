#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Template: Voyage-style travel itinerary DOCX.

1. Set TEMPLATE to a known-good Voyage itinerary (keeps header/logo).
2. Fill TITLE / SUBTITLE / DAYS / PRICING from the supplier quote only.
3. Run:
       PYTHONIOENCODING=utf-8 python scripts/generate.py

Do NOT invent sights, meals, hotels, gifts, or prices.
See references/fact-check.md before delivery.

Worked example: generate_chaoshan6_voyage_style.py
"""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# -- CONFIG (edit per quote) -------------------------------------------------

TEMPLATE = r'C:\Users\DELL\Downloads\12D11N_东北草原深度游_2026年6月.docx'

# Primary output + optional ASCII copies (same folder as source .doc preferred)
OUTPUTS = [
    r'C:\Users\DELL\Downloads\itinerary_output.docx',
    r'C:\Users\DELL\Downloads\itinerary_output_en.docx',
]

TITLE = '行程标题（来自报价）'
SUBTITLE = '城市A·城市B·城市C'
NOTE_NO_SHOP = '***全程不进店'  # or quote shopping sentence; leave '' if none

HIGHLIGHT_LINES = [
    '亮点一句，必须来自报价',
]

FOOD = '各地美食列表，来自报价'
GIFTS = '赠送项，仅报价明确为赠送者'

# Each day: quote facts only
DAYS = [
    {
        'day': 'Day 1',
        'route': '出发地→目的地',
        'meals': '晚：示例餐',
        'summary': '当日行程概述，紧贴报价原文。',
        'sights': [
            ('景点名', '短说明，不新增事实。'),
        ],
        'hotel': '酒店名（五星）',  # or None on return day
    },
]

PRICING = {
    'main': '报价：RMB0/人，单间差RMB0，……（整句抄自报价）',
    'shopping': '全程不进店',
    'tip': 'RMB0/人天',
    'gift_summary': '每人每天矿泉水两瓶',
    'note': '备注：有效期……',
    'hotels': '精选酒店：……',
    'meals': '膳食（据报价）：……',
    'contacts': [
        '联系人：……',
        '电话：……',
        'Email：……',
    ],
}

# -- COLORS (Voyage / 东北草原) -----------------------------------------------

C_TITLE = RGBColor(0x1A, 0x6B, 0x5E)
C_SUB = RGBColor(0x55, 0x55, 0x55)
C_ORANGE = RGBColor(0xC4, 0x59, 0x11)
C_DAY_BG = '2E8B7A'
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY = RGBColor(0x00, 0x00, 0x00)
C_HOTEL = RGBColor(0x1A, 0x6B, 0x5E)

FONT_CN = 'SimSun'
FONT_DAY = 'Arial Unicode MS'


# -- HELPERS ------------------------------------------------------------------

def set_run_font(run, name=FONT_CN, size_pt=11, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    # Never set bold=False — Word rejects w:b w:val="0"
    if bold is True:
        run.font.bold = True
        if rPr.find(qn('w:bCs')) is None:
            rPr.append(OxmlElement('w:bCs'))
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def set_spacing(paragraph, before=None, after=None, left=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if left is not None:
        pf.left_indent = Pt(left)
    if line is not None:
        pf.line_spacing = line


def clear_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def add_para(doc, text, *, size=11, bold=False, color=C_BODY, font=FONT_CN,
             align=None, before=0, after=4, left=None, justify=False,
             shade=None, line=1.15):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align is not None:
        p.alignment = align
    set_spacing(p, before=before, after=after, left=left, line=line)
    if shade:
        set_paragraph_shading(p, shade)
    if text:
        run = p.add_run(text)
        set_run_font(
            run, name=font, size_pt=size,
            bold=True if bold else None, color=color,
        )
    return p


def add_sight_para(doc, name: str, desc: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=2, after=3, line=1.2)
    r1 = p.add_run(f'【{name}】')
    set_run_font(r1, size_pt=10.5, bold=True, color=C_TITLE)
    r2 = p.add_run(desc)
    set_run_font(r2, size_pt=10.5, color=C_BODY)
    return p


def add_label_value(doc, label: str, value: str, *, before=2, after=2):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after, line=1.15)
    r1 = p.add_run(label)
    set_run_font(r1, size_pt=10.5, bold=True, color=C_TITLE)
    r2 = p.add_run(value)
    set_run_font(r2, size_pt=10.5, color=C_BODY)
    return p


# -- BUILD --------------------------------------------------------------------

def build():
    if not os.path.isfile(TEMPLATE):
        raise SystemExit(
            f'Template not found:\n  {TEMPLATE}\n'
            'Set TEMPLATE to a known-good Voyage itinerary .docx'
        )

    doc = Document(TEMPLATE)
    clear_body(doc)

    add_para(doc, TITLE, size=18, bold=True, color=C_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=3, line=1.2)
    add_para(doc, SUBTITLE, size=11, bold=True, color=C_SUB,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line=1.15)
    if NOTE_NO_SHOP:
        add_para(doc, NOTE_NO_SHOP, size=12, bold=True, color=C_ORANGE,
                 align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8, line=1.1)

    if HIGHLIGHT_LINES:
        add_label_value(
            doc, '独家安排：',
            '；'.join(HIGHLIGHT_LINES) + '。',
            before=2, after=4,
        )
    if FOOD:
        add_label_value(doc, '各地美食：', FOOD.rstrip('。') + '。', before=1, after=3)
    if GIFTS:
        add_label_value(doc, '赠送：', GIFTS.rstrip('。') + '。', before=1, after=8)

    for day in DAYS:
        # No giant space padding — 2–4 spaces only
        header = f'{day["day"]}  {day["route"]}    （{day["meals"]}）'
        add_para(doc, header, size=11, bold=True, color=C_WHITE, font=FONT_DAY,
                 before=10, after=4, left=2, shade=C_DAY_BG, line=1.1)

        add_para(doc, day['summary'], size=10.5, color=C_BODY, justify=True,
                 before=3, after=3, line=1.25)

        for name, desc in day.get('sights') or []:
            add_sight_para(doc, name, desc)

        if day.get('hotel'):
            add_para(doc, f'宿：{day["hotel"]}', size=10.5, bold=True, color=C_HOTEL,
                     before=4, after=2, left=2, line=1.1)

    add_para(doc, '报价与说明', size=12, bold=True, color=C_TITLE,
             before=12, after=6, line=1.1)

    add_para(doc, PRICING['main'], size=10.5, bold=True, color=C_BODY,
             before=2, after=4, line=1.2)

    if PRICING.get('shopping'):
        add_label_value(doc, '购物：', PRICING['shopping'], before=1, after=1)
    if PRICING.get('tip'):
        add_label_value(doc, '小费：', PRICING['tip'], before=0, after=1)
    if PRICING.get('gift_summary'):
        add_label_value(doc, '赠送：', PRICING['gift_summary'], before=0, after=3)
    if PRICING.get('note'):
        add_para(doc, PRICING['note'], size=10.5, color=C_BODY,
                 before=2, after=4, line=1.2)
    if PRICING.get('hotels'):
        add_para(doc, PRICING['hotels'], size=10.5, color=C_BODY,
                 before=1, after=4, line=1.2)
    if PRICING.get('meals'):
        add_para(doc, PRICING['meals'], size=10.5, color=C_BODY,
                 before=1, after=6, line=1.25)

    add_para(doc, '联系方式', size=12, bold=True, color=C_TITLE,
             before=4, after=3, line=1.1)
    for line in PRICING.get('contacts') or []:
        add_para(doc, line, size=10.5, color=C_BODY, before=0, after=1, line=1.15)

    first = OUTPUTS[0]
    os.makedirs(os.path.dirname(first) or '.', exist_ok=True)
    doc.save(first)
    print('Saved:', first, os.path.getsize(first))
    for other in OUTPUTS[1:]:
        shutil.copy2(first, other)
        print('Copied:', other)


if __name__ == '__main__':
    build()
